from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    r_text = OxmlElement('w:t')
    r_text.text = "1"
    run._r.append(r_text)
    run._r.append(fldChar3)

doc = Document('/Users/miteshsingh/Downloads/miteshsynopsis.docx')

for i, section in enumerate(doc.sections):
    footer = section.footer
    if i > 0:
        footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    add_page_number(run)

    sectPr = section._sectPr
    pgNumType = OxmlElement('w:pgNumType')
    if i == 0:
        # Cover page starts at 1, TOC is page 2-3, content starts page 3
        pgNumType.set(qn('w:fmt'), 'decimal')
        pgNumType.set(qn('w:start'), '1')
    sectPr.append(pgNumType)

doc.save('/Users/miteshsingh/Downloads/miteshsynopsis.docx')
print("Synopsis footers applied!")
