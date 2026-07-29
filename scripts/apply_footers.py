import sys
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    # The rendered value
    r_text = OxmlElement('w:t')
    r_text.text = "1"
    run._r.append(r_text)
    run._r.append(fldChar3)

doc = Document('/Users/miteshsingh/Downloads/miteshreport.docx')

# Remove manual fake page numbers that might be floating
for p in doc.paragraphs:
    txt = p.text.strip()
    if txt in ["i", "ii", "iii", "iv", "v", "vi", "vii", "1", "2", "3", "4", "5", "6", "7", "8", "9", "10", "11", "12", "13", "14", "15", "16"]:
        if p.alignment == WD_ALIGN_PARAGRAPH.CENTER and len(txt) <= 3:
            p.clear()

# Add section breaks before Chapter 1 to restart numbering
# Since docx cannot easily insert section breaks in middle, we rely on the footers
for i, section in enumerate(doc.sections):
    footer = section.footer
    if i > 0:
        footer.is_linked_to_previous = False
        
    for p in footer.paragraphs:
        p.clear()
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    add_page_number(run)

    # Set formatting based on section
    sectPr = section._sectPr
    pgNumType = OxmlElement('w:pgNumType')
    if i == 0:
        pgNumType.set(qn('w:fmt'), 'lowerRoman')
        pgNumType.set(qn('w:start'), '1')
    else:
        pgNumType.set(qn('w:fmt'), 'decimal')
        pgNumType.set(qn('w:start'), '1')
    sectPr.append(pgNumType)

doc.save('/Users/miteshsingh/Downloads/miteshreport.docx')
print("Footers applied!")
