import os
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('/Users/miteshsingh/Downloads/report mitsh.docx')
IMG_DIR = "/Users/miteshsingh/Documents/projects/chess_analyzer/orivis_images"
DATA_DIR = "/Users/miteshsingh/Documents/projects/chess_analyzer/data/analytics"

# Find the paragraph that says "Figure 4.1" and the one before it (the image placeholder)
# We need to find where Chapter 4.3 "Flow Chart / Activity Diagram" content is and replace with proper figures

# Strategy: find paragraphs by text content and insert images after them

def find_para_index(doc, text_contains):
    for i, p in enumerate(doc.paragraphs):
        if text_contains in p.text:
            return i
    return -1

def insert_image_after(doc, para_index, img_path, caption_text, width=Inches(4.5)):
    """Insert an image and caption after a given paragraph index"""
    if not os.path.exists(img_path):
        print(f"WARNING: {img_path} not found!")
        return
    
    # Add image
    p_img = doc.paragraphs[para_index]
    # We'll add a new paragraph after with the image
    # Actually, let's just add to the end since we're rebuilding Chapter 4 and 5
    
    run = p_img.add_run()
    run.add_picture(img_path, width=width)
    
    print(f"Added image: {caption_text}")

# Let's find key paragraphs
for i, p in enumerate(doc.paragraphs):
    if "Activity Diagram and Dashboard Flow" in p.text:
        print(f"Found Figure 4.1 caption at para {i}")
    if "Blunder Analysis Result" in p.text:
        print(f"Found Figure 5.1 caption at para {i}")
    if "Flow Chart / Activity Diagram" in p.text and "4.3" in p.text:
        print(f"Found Section 4.3 heading at para {i}")

print("\nTotal paragraphs:", len(doc.paragraphs))
print("Total inline shapes (images):", len(doc.inline_shapes))

doc.save('/Users/miteshsingh/Downloads/report mitsh.docx')
print("Debug done")
